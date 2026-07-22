import os
import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document


class IndustrialLoader:

    def __init__(self):
        self.base_path = "data"

    def read_pdf(self, path):
        text = ""

        try:
            reader = PdfReader(path)

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        except Exception as e:
            print(f"Error reading PDF {path}: {e}")

        return text.strip()

    def read_docx(self, path):
        doc = DocxDocument(path)
        text = "\n".join([p.text for p in doc.paragraphs])

        return text.strip()

    def load_documents(self):

        documents = []

        folders = {
            "manuals": [".pdf"],
            "sops": [".pdf", ".docx"],
            "regulations": [".pdf"]
        }

        # ----------------------------
        # PDF / DOCX
        # ----------------------------

        for folder, extensions in folders.items():

            folder_path = os.path.join(self.base_path, folder)

            if not os.path.exists(folder_path):
                continue

            for file in os.listdir(folder_path):

                if not any(file.lower().endswith(ext) for ext in extensions):
                    continue

                full_path = os.path.join(folder_path, file)

                if file.lower().endswith(".pdf"):
                    text = self.read_pdf(full_path)
                else:
                    text = self.read_docx(full_path)

                # Skip empty files
                if len(text.strip()) == 0:
                    print(f"Skipped empty file: {file}")
                    continue

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file,
                            "category": folder
                        }
                    )
                )

        # ----------------------------
        # CSV
        # ----------------------------

        csv_folders = [
            "maintenance_logs",
            "inspection_reports",
            "incident_reports"
        ]

        for folder in csv_folders:

            folder_path = os.path.join(self.base_path, folder)

            if not os.path.exists(folder_path):
                continue

            for file in os.listdir(folder_path):

                if not file.endswith(".csv"):
                    continue

                full_path = os.path.join(folder_path, file)

                df = pd.read_csv(full_path)

                # One document per row
                for idx, row in df.iterrows():

                    content = "\n".join(
                        [f"{col}: {row[col]}" for col in df.columns]
                    )

                    documents.append(
                        Document(
                            page_content=content,
                            metadata={
                                "source": file,
                                "category": folder,
                                "row": idx
                            }
                        )
                    )

        return documents


if __name__ == "__main__":

    loader = IndustrialLoader()

    docs = loader.load_documents()

    print("=" * 70)
    print("TOTAL DOCUMENTS :", len(docs))
    print("=" * 70)

    for doc in docs:
        print(doc.metadata)